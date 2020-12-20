# -*- coding: utf-8 -*-
"""
Created on Tue Sep 15 12:39:37 2020

@author: zhen chen

MIT Licence.

Python version: 3.7


Description: read doc files by python
    
"""


import docx  # docx 文件可以直接读取，读取 07版本之前的 doc文件需要安装 win32com 包，安装方法：conda install -c conda-forge python-docx
import win32com.client as wc # 将 doc 转化为 docx，安装方法：conda install -c anaconda pywin32
import os # 读取文件的包
import re # 正则表达式包，处理字符串
import pandas as pd


# 读取文件夹所有文件的文件名
path = 'E:\\爬虫练习\\word文件'  #文件夹目录
files = os.listdir(path) #得到文件夹下的所有文件名称，不要文件夹名


# 将 doc文件另存为 docx
word = wc.Dispatch("Word.Application")
for file in files:    
    path_file = path + '\\' + file # 一般要用 \\，防止 python 转义 \
    doc = word.Documents.Open(path_file)
    new_path = 'E:\\爬虫练习' + '\\docx 文件夹\\' # win32包读文件一般地址只能用“\\”，不能用“/”，哪怕加了 r 也不行，涉及到将反斜杠看成转义字符  
    isExists=os.path.exists(new_path) # 判断新文件夹路径是否存在
    #若不存在则创建文件夹
    if not isExists:
        os.makedirs(new_path) # 创建新文件夹
    new_path_file = new_path + file[:-3] + 'docx' 
    if not os.path.exists(new_path_file): # 若没有改文件名则新建文件
        doc.SaveAs(new_path_file, 16) #注意 SaveAs会打开保存后的文件，有时可能看不到，但后台一定是打开的, 处理 word文件后面跟数字 16
    doc.Close
word.Quit 


files= os.listdir(new_path) #得到文件夹下的所有文件名称    
# 读取日期与部分数字
dates = []
number = []
for file in files:
    file = docx.Document(new_path + file) # 用 docx 读取
    for p in file.paragraphs:
         # 使用正则表达式获取文本信息
         if ('审批业务' in p.text):
             string = p.text
             results = re.search('.*?(\d+年\d+月\d+日).*?共有(\d+).*?', string, re.S)  # 评论者账号, re.S 表示可以换行查找
             dates.append(results.group(1))     
             number.append(results.group(2))   

df = pd.DataFrame({'日期':dates, '核查项目数':number})
df.to_excel(r'E:\record.xlsx')


# 读取联系方式       
#names = []
#phones = []
#emails = []          
#for file in files:
#    file = docx.Document(new_path + file) # 用 docx 读取
#    for p in file.paragraphs:
#         # 使用正则表达式获取文本信息
#         if ('联系人' in p.text):
#             string = p.text
#             # () 表示想要输出的内容，.*表示匹配任意字符或表达式
#             results = re.search('联系人：(\D\S+).*?(1\d{10}).*电子邮箱：(.*)', string, re.S)  # 评论者账号, re.S 表示可以换行查找
#             print(results.group(1))
#             names.append(results.group(1))
#             phones.append(results.group(2))
#             emails.append(results.group(3))






  